#!/usr/bin/env python3
"""Construye el documento Word del informe de resultados."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from _bootstrap import setup

setup()

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt

from generate_entregable2_docx import (
    UCR_BLUE,
    TEXT_DARK,
    TEXT_MUTED,
    add_body,
    add_table,
    configure_document,
)
from informe_data_local import AnalysisBundle
from informe_figures_local import (
    caption,
    FIG_CHAT_EXCHANGES,
    FIG_CHAT_HELPSCORE,
    FIG_CHAT_LEAKS,
    FIG_MECUE,
    FIG_PERFIL,
    FIG_RAW_TLX,
    FIG_RQ1_PRECISION,
    FIG_RQ2_CONFIDENCE,
    FIG_RQ3_CURVE_BY_ITEM,
    FIG_RQ3_GAP,
    FIG_RQ3_GAP_BY_ITEM,
    FIG_TIME,
)
from informe_narrative_local import InformeNarrative
from informe_synthesis_local import (
    build_flujo_summary_rows,
    build_sessions_detail_rows,
    build_viabilidad_semaforo_rows,
)

AUTHOR = "Ney Fred Jiménez (B03230)"
COURSE = "PF-3311 — Agentes Virtuales Inteligentes"
STUDY_TITLE = (
    "Efecto de la visibilidad de un agente virtual en el desempeño "
    "y la confianza del usuario en tareas de decisión basadas en reglas"
)


def _add_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = UCR_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)
    r2.font.color.rgb = TEXT_DARK

    for line in (COURSE, AUTHOR, f"Generado: {datetime.now().strftime('%d/%m/%Y')}"):
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(line)
        r3.font.size = Pt(11)
        r3.font.color.rgb = TEXT_MUTED

    doc.add_page_break()


def _csv_to_table_rows(
    rows: list[dict[str, str]],
    columns: list[tuple[str, str]],
) -> list[list[str]]:
    table_rows = [[label for label, _ in columns]]
    for row in rows:
        table_rows.append([row.get(key, "") for _, key in columns])
    return table_rows


def _add_csv_table(doc: Document, title: str, rows: list[dict[str, str]], columns: list[tuple[str, str]]) -> None:
    if not rows:
        add_body(doc, f"(Sin datos para {title})")
        return
    doc.add_heading(title, level=2)
    add_table(doc, _csv_to_table_rows(rows, columns))


def _add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.is_file():
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = TEXT_MUTED
    doc.add_paragraph()


def _add_interpretation(doc: Document, section_key: str, narrative: InformeNarrative) -> None:
    paragraphs = narrative.sections.get(section_key, [])
    if not paragraphs:
        return
    doc.add_heading("Lectura de la figura", level=2)
    for para in paragraphs:
        add_body(doc, para)


def _add_glossary(doc: Document) -> None:
    doc.add_heading("Guía rápida de términos", level=2)
    for line in (
        "A / B / C: sin agente, chat de texto, avatar con voz.",
        "Precisión: % de respuestas correctas en el bloque (6 preguntas).",
        "Confianza: qué tan seguro se sintió el participante (1 = nada, 7 = mucho).",
        "Brecha de calibración: desajuste entre confianza y aciertos; positivo ≈ más confianza que aciertos.",
        "HelpScore: utilidad pedagógica del chat (evaluación automática).",
        "RAW-TLX: carga de trabajo percibida tras cada bloque (más alto = más exigente).",
        "meCUE: experiencia con el agente conversacional (bloques B y C); Mód. II solo en C.",
        "Tiempo de respuesta: segundos entre ver la pregunta y pulsar entregar (Unity).",
    ):
        add_body(doc, line)


def _add_dict_table(
    doc: Document,
    title: str,
    rows: list[dict[str, str]],
    columns: list[tuple[str, str]],
) -> None:
    if not rows:
        return
    doc.add_heading(title, level=2)
    table_rows = [[label for label, _ in columns]]
    for row in rows:
        table_rows.append([row.get(key, "") for _, key in columns])
    add_table(doc, table_rows)


def _add_descriptive_table(
    doc: Document,
    title: str,
    rows: list[dict[str, str]],
) -> None:
    if not rows:
        return
    _add_csv_table(
        doc,
        title,
        rows,
        [
            ("Cond.", "Condition"),
            ("Métrica", "Metric"),
            ("N", "N"),
            ("Media", "Mean"),
            ("Mediana", "Median"),
            ("DE", "StDev"),
            ("Mín.", "Min"),
            ("Máx.", "Max"),
        ],
    )


def _add_section_paragraphs(doc: Document, heading: str, paragraphs: list[str]) -> None:
    doc.add_heading(heading, level=1)
    for para in paragraphs:
        add_body(doc, para)


def build_informe_docx(
    bundle: AnalysisBundle,
    narrative: InformeNarrative,
    output_path: Path,
    title: str = "Informe de resultados — PF-3311",
) -> Path:
    doc = Document()
    configure_document(doc)

    _add_cover(doc, title, STUDY_TITLE)

    doc.add_heading("Índice", level=1)
    for item in (
        "1. Resumen y guía de términos",
        "2. Muestra, flujo y metodología",
        "3. RQ1: Precisión",
        "4. RQ2: Confianza",
        "5. RQ3: Calibración",
        "6. Cuestionarios (RAW-TLX / meCUE)",
        "7. Agente (B vs C)",
        "8. Viabilidad técnica",
        "9. Respuestas a las preguntas de investigación",
        "10. Discusión",
        "11. Conclusiones",
        "Referencias",
        "Anexo: Inferencia estadística",
    ):
        add_body(doc, item)
    doc.add_page_break()

    _add_section_paragraphs(doc, "1. Resumen", narrative.sections.get("resumen", []))
    _add_glossary(doc)

    doc.add_heading("2. Muestra, flujo y metodología", level=1)
    for para in narrative.sections.get("muestra", []):
        add_body(doc, para)
    for para in narrative.sections.get("flujo", []):
        add_body(doc, para)
    _add_dict_table(
        doc,
        "Tabla F. Embudo de participantes",
        build_flujo_summary_rows(bundle),
        [
            ("Etapa", "Etapa"),
            ("N", "N"),
            ("Nota", "Nota"),
        ],
    )
    _add_dict_table(
        doc,
        "Tabla S. Detalle por participante canónico",
        build_sessions_detail_rows(bundle),
        [
            ("P##", "P##"),
            ("Ítems A", "Ítems A"),
            ("Ítems B", "Ítems B"),
            ("Ítems C", "Ítems C"),
            ("Completa", "Completa"),
            ("Consent.", "Consent."),
        ],
    )
    for para in narrative.sections.get("metodologia", []):
        add_body(doc, para)

    for para in narrative.sections.get("perfil", []):
        add_body(doc, para)
    if bundle.perfil:
        _add_csv_table(
            doc,
            "Tabla P. Perfil de participantes (Form 0)",
            bundle.perfil,
            [
                ("P##", "ParticipantCode"),
                ("Edad", "AgeRange"),
                ("Educación", "Education"),
                ("Asistentes digitales", "AssistantFrequency"),
                ("Avatares/agentes", "AvatarExperience"),
            ],
        )
        _add_figure(
            doc,
            bundle.figures_dir / "perfil_muestra.png",
            caption(FIG_PERFIL, "Perfil de la muestra (Formulario 0)"),
        )
        _add_interpretation(doc, "perfil_figura_interpretacion", narrative)

    doc.add_heading("3. RQ1: Precisión", level=1)
    for para in narrative.sections.get("rq1", []):
        add_body(doc, para)
    _add_csv_table(
        doc,
        "Tabla 1. Precisión por condición (% aciertos)",
        bundle.rq1_group,
        [("Condición", "Condition"), ("Media %", "MeanAccuracyPct"), ("N", "NParticipants")],
    )
    _add_descriptive_table(doc, "Descriptivos de precisión", bundle.rq1_desc)
    _add_csv_table(
        doc,
        "Precisión por participante",
        bundle.rq1_participant,
        [
            ("Participante", "ParticipantSession"),
            ("A %", "AccuracyPct_A"),
            ("B %", "AccuracyPct_B"),
            ("C %", "AccuracyPct_C"),
        ],
    )
    _add_figure(
        doc,
        bundle.figures_dir / "rq1_precision.png",
        caption(FIG_RQ1_PRECISION, "Precisión media por condición"),
    )
    _add_interpretation(doc, "rq1_interpretacion", narrative)

    doc.add_heading("4. RQ2: Confianza", level=1)
    for para in narrative.sections.get("rq2_unity", []):
        add_body(doc, para)
    _add_csv_table(
        doc,
        "Tabla 2. Confianza media (Unity, 1–7)",
        bundle.rq2_group,
        [("Condición", "Condition"), ("Media", "MeanConfidence"), ("N", "NParticipants")],
    )
    _add_descriptive_table(doc, "Descriptivos de confianza", bundle.rq2_desc)
    _add_figure(
        doc,
        bundle.figures_dir / "rq2_confidence.png",
        caption(FIG_RQ2_CONFIDENCE, "Confianza media por condición"),
    )
    _add_interpretation(doc, "rq2_interpretacion", narrative)

    doc.add_heading("5. RQ3: Calibración", level=1)
    for para in narrative.sections.get("rq3", []):
        add_body(doc, para)
    _add_csv_table(
        doc,
        "Tabla 3. Brecha de calibración por condición",
        bundle.rq3_group,
        [("Condición", "Condition"), ("Brecha media", "MeanCalibrationGap"), ("N", "NParticipants")],
    )
    _add_descriptive_table(doc, "Descriptivos de brecha de calibración", bundle.rq3_desc)
    _add_figure(
        doc,
        bundle.figures_dir / "rq3_calibration_gap.png",
        caption(FIG_RQ3_GAP, "Brecha de calibración por condición"),
    )
    _add_interpretation(doc, "rq3_interpretacion", narrative)
    if bundle.rq3_calibration_by_item:
        for para in narrative.sections.get("rq3_items", []):
            add_body(doc, para)
        _add_csv_table(
            doc,
            "Tabla 3b. Calibración por ítem (Q1–Q6)",
            bundle.rq3_calibration_by_item,
            [
                ("Condición", "Condition"),
                ("Ítem", "QuestionNumber"),
                ("N", "N"),
                ("Conf. norm.", "MeanConfidenceNorm"),
                ("% aciertos", "AccuracyPct"),
                ("Brecha", "CalibrationGap"),
            ],
        )
        _add_figure(
            doc,
            bundle.figures_dir / "rq3_calibration_gap_by_item.png",
            caption(FIG_RQ3_GAP_BY_ITEM, "Brecha de calibración por ítem"),
        )
        _add_interpretation(doc, "rq3_fig3b_interpretacion", narrative)
        _add_figure(
            doc,
            bundle.figures_dir / "rq3_calibration_curve_by_item.png",
            caption(FIG_RQ3_CURVE_BY_ITEM, "Curva de calibración por ítem"),
        )
        _add_interpretation(doc, "rq3_fig3c_interpretacion", narrative)

    doc.add_heading("Evaluación de hipótesis H1–H3", level=2)
    for para in narrative.sections.get("hipotesis", []):
        add_body(doc, para)

    doc.add_heading("6. Cuestionarios (RAW-TLX / meCUE)", level=1)
    for para in narrative.sections.get("rq2_forms", []):
        add_body(doc, para)
    if bundle.rq2_forms_tlx_group:
        _add_csv_table(
            doc,
            "Tabla 4. RAW-TLX por condición",
            bundle.rq2_forms_tlx_group,
            [("Condición", "Condition"), ("Media", "MeanRAWTLX"), ("N", "NParticipants")],
        )
    if bundle.rq2_forms_tlx:
        _add_csv_table(
            doc,
            "RAW-TLX por participante",
            bundle.rq2_forms_tlx,
            [
                ("P##", "ParticipantCode"),
                ("A", "RAW_TLX_A"),
                ("B", "RAW_TLX_B"),
                ("C", "RAW_TLX_C"),
            ],
        )
    _add_figure(
        doc,
        bundle.figures_dir / "forms_raw_tlx_by_condition.png",
        caption(FIG_RAW_TLX, "Carga de trabajo (RAW-TLX) por condición"),
    )
    _add_interpretation(doc, "forms_tlx_interpretacion", narrative)

    if bundle.rq2_forms_mecue_group:
        _add_csv_table(
            doc,
            "Tabla 5. meCUE: medias B vs C",
            bundle.rq2_forms_mecue_group,
            [
                ("Módulo", "ModuleLabel"),
                ("Media B", "Mean_B"),
                ("N B", "N_B"),
                ("Media C", "Mean_C"),
                ("N C", "N_C"),
            ],
        )
    if bundle.rq2_forms_mecue:
        _add_csv_table(
            doc,
            "meCUE por participante (B vs C)",
            bundle.rq2_forms_mecue,
            [
                ("P##", "ParticipantCode"),
                ("Módulo", "Module"),
                ("B", "Score_B"),
                ("C", "Score_C"),
            ],
        )
    _add_figure(
        doc,
        bundle.figures_dir / "forms_mecue_b_vs_c.png",
        caption(FIG_MECUE, "meCUE: comparación B vs C por subescala"),
    )
    _add_interpretation(doc, "mecue_interpretacion", narrative)

    for para in narrative.sections.get("mecue_ii", []):
        add_body(doc, para)
    if bundle.rq2_forms_mecue_ii_group:
        _add_csv_table(
            doc,
            "Tabla 5b. meCUE Módulo II (solo condición C)",
            bundle.rq2_forms_mecue_ii_group,
            [
                ("Módulo", "Module"),
                ("Media C", "Mean_C"),
                ("N", "NParticipants"),
            ],
        )
    if bundle.rq2_forms_mecue_ii_participant:
        _add_csv_table(
            doc,
            "meCUE Módulo II por participante",
            bundle.rq2_forms_mecue_ii_participant,
            [
                ("P##", "ParticipantCode"),
                ("Puntuación C", "MECUE_II_C"),
            ],
        )

    doc.add_heading("6.1 Tiempo de respuesta", level=2)
    for para in narrative.sections.get("tiempo", []):
        add_body(doc, para)
    if bundle.items_time_by_condition:
        _add_csv_table(
            doc,
            "Tabla 6. Tiempo medio por condición",
            bundle.items_time_by_condition,
            [("Condición", "Condition"), ("Segundos", "MeanTimeSeconds"), ("Ítems", "NItems")],
        )
    if bundle.items_time_spent:
        _add_csv_table(
            doc,
            "Tiempo por ítem",
            bundle.items_time_spent,
            [
                ("Cond.", "Condition"),
                ("Esc.", "ScenarioNumber"),
                ("Preg.", "QuestionNumber"),
                ("N", "N"),
                ("s", "MeanTimeSeconds"),
            ],
        )
    _add_figure(
        doc,
        bundle.figures_dir / "items_time_by_condition.png",
        caption(FIG_TIME, "Tiempo medio de respuesta por condición"),
    )
    _add_interpretation(doc, "tiempo_interpretacion", narrative)

    doc.add_heading("7. Agente (B vs C)", level=1)
    for para in narrative.sections.get("agente", []):
        add_body(doc, para)
    _add_csv_table(
        doc,
        "Métricas agregadas B vs C",
        bundle.chat_group,
        [("Métrica", "Metric"), ("Media B", "Mean_B"), ("N B", "N_B"), ("Media C", "Mean_C"), ("N C", "N_C")],
    )
    for fig_name, fig_num, fig_title, interp_key in (
        ("chat_helpscore_b_vs_c.png", FIG_CHAT_HELPSCORE, "HelpScore B vs C", "agente_fig4_interpretacion"),
        ("chat_exchanges_b_vs_c.png", FIG_CHAT_EXCHANGES, "Intercambios de chat B vs C", "agente_fig5_interpretacion"),
        ("chat_leaks_b_vs_c.png", FIG_CHAT_LEAKS, "Leaks del modelo B vs C", "agente_fig6_interpretacion"),
    ):
        _add_figure(doc, bundle.figures_dir / fig_name, caption(fig_num, fig_title))
        _add_interpretation(doc, interp_key, narrative)

    doc.add_heading("8. Viabilidad técnica", level=1)
    for para in narrative.sections.get("viabilidad", []):
        add_body(doc, para)
    if bundle.pilot_integrity:
        _add_csv_table(
            doc,
            "Resumen de viabilidad",
            bundle.pilot_integrity,
            [("Métrica", "Metric"), ("Valor", "Value")],
        )
    _add_dict_table(
        doc,
        "Tabla V. Semáforo de viabilidad",
        build_viabilidad_semaforo_rows(bundle),
        [
            ("Criterio", "Criterio"),
            ("Meta del diseño", "Meta del diseño"),
            ("Valor observado", "Valor observado"),
            ("Estado", "Estado"),
        ],
    )

    _add_section_paragraphs(
        doc,
        "9. Respuestas a las preguntas de investigación",
        narrative.sections.get("respuestas_rq", []),
    )
    _add_section_paragraphs(doc, "10. Discusión", narrative.sections.get("discusion", []))
    _add_section_paragraphs(doc, "11. Conclusiones", narrative.sections.get("conclusiones", []))

    doc.add_heading("Referencias", level=1)
    for ref in narrative.sections.get("referencias", []):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.add_run(ref)

    doc.add_page_break()
    doc.add_heading("Anexo: Inferencia estadística", level=1)
    add_body(doc, "Registro de pruebas Friedman y Wilcoxon aplicadas al generar este informe.")
    if bundle.inference_omnibus:
        _add_csv_table(
            doc,
            "Pruebas omnibus (Friedman)",
            bundle.inference_omnibus,
            [
                ("Análisis", "AnalysisId"),
                ("Métrica", "Metric"),
                ("Etiqueta", "Label"),
                ("Prueba", "Test"),
                ("Estadístico", "Statistic"),
                ("p", "PValue"),
                ("Kendall W", "KendallW"),
                ("N", "N"),
            ],
        )
    if bundle.inference_pairwise:
        _add_csv_table(
            doc,
            "Post hoc pareado (Wilcoxon)",
            bundle.inference_pairwise,
            [
                ("Análisis", "AnalysisId"),
                ("Contraste", "Left"),
                ("vs", "Right"),
                ("W", "Statistic"),
                ("p", "PValue"),
                ("p adj. Bonf.", "PAdjustedBonferroni"),
                ("p adj. Holm", "PAdjustedHolm"),
            ],
        )
    if bundle.inference_bootstrap:
        _add_csv_table(
            doc,
            "IC bootstrap 95 % (diferencia pareada)",
            bundle.inference_bootstrap,
            [
                ("Análisis", "AnalysisId"),
                ("Contraste", "Left"),
                ("vs", "Right"),
                ("Δ media", "MeanDiff"),
                ("IC inf.", "CI_Low"),
                ("IC sup.", "CI_High"),
                ("Unidades", "Units"),
            ],
        )
    if bundle.inference_contrasts:
        _add_csv_table(
            doc,
            "Contrastes dirigidos (H3, B vs C, HelpScore)",
            bundle.inference_contrasts,
            [
                ("ID", "ContrastId"),
                ("Descripción", "Description"),
                ("Δ", "MeanDiff"),
                ("IC inf.", "CI_Low"),
                ("IC sup.", "CI_High"),
                ("p", "PValue"),
                ("N", "N"),
            ],
        )
    for line in narrative.inference_log:
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
