#!/usr/bin/env python3
"""Construye artículo tipo paper en Word."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

from _bootstrap import setup

setup()

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from generate_entregable2_docx import (
    TEXT_DARK,
    TEXT_MUTED,
    UCR_BLUE,
    add_body,
    add_table,
    configure_document,
)
from informe_data_local import AnalysisBundle
from informe_docx_local import _csv_to_table_rows
from paper_narrative_local import PaperNarrative

COLUMN_GAP_TWIPS = 720  # ~0,5 pulgada entre columnas


def _set_columns(section, num: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    if num > 1:
        cols.set(qn("w:space"), str(COLUMN_GAP_TWIPS))
    elif qn("w:space") in cols.attrib:
        del cols.attrib[qn("w:space")]


def _begin_two_column_body(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_columns(section, 2)


def _with_full_width(doc: Document, builder: Callable[[Document], None]) -> None:
    """Inserta bloque a ancho completo (tabla/figura) entre secciones de dos columnas."""
    single = doc.add_section(WD_SECTION_START.CONTINUOUS)
    _set_columns(single, 1)
    builder(doc)
    double = doc.add_section(WD_SECTION_START.CONTINUOUS)
    _set_columns(double, 2)


def _add_centered(doc: Document, text: str, *, size: int = 11, bold: bool = False, color=TEXT_DARK) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _add_section(doc: Document, title: str, paragraphs: list[str]) -> None:
    doc.add_heading(title, level=1)
    for para in paragraphs:
        add_body(doc, para)


def _add_csv_table(doc: Document, title: str, rows: list[dict[str, str]], columns: list[tuple[str, str]]) -> None:
    if not rows:
        add_body(doc, f"(Sin datos para {title})")
        return
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    add_table(doc, _csv_to_table_rows(rows, columns))


def _add_figure(doc: Document, path: Path, caption: str, *, width_inches: float = 6.2) -> None:
    if not path.is_file():
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = TEXT_MUTED
    doc.add_paragraph()


def build_paper_docx(
    bundle: AnalysisBundle,
    paper: PaperNarrative,
    output_path: Path,
) -> Path:
    doc = Document()
    configure_document(doc)
    _set_columns(doc.sections[0], 1)

    _add_centered(doc, paper.title, size=16, bold=True, color=UCR_BLUE)
    _add_centered(doc, paper.authors, size=12, bold=True)
    _add_centered(doc, paper.affiliation, size=10, color=TEXT_MUTED)
    doc.add_paragraph()

    doc.add_heading("Resumen", level=2)
    for para in paper.abstract:
        add_body(doc, para)
    add_body(doc, f"**Palabras clave:** {paper.keywords}")

    _begin_two_column_body(doc)

    _add_section(doc, "1. Introducción", paper.sections.get("introduccion", []))
    _add_section(doc, "2. Método", paper.sections.get("metodo", []))
    if bundle.perfil:
        _with_full_width(
            doc,
            lambda d: _add_csv_table(
                d,
                "Tabla 0. Perfil de participantes (Form 0)",
                bundle.perfil,
                [
                    ("P##", "ParticipantCode"),
                    ("Edad", "AgeRange"),
                    ("Educación", "Education"),
                    ("Asistentes digitales", "AssistantFrequency"),
                    ("Avatares/agentes", "AvatarExperience"),
                ],
            ),
        )
    _add_section(doc, "3. Resultados", paper.sections.get("resultados", []))

    _with_full_width(
        doc,
        lambda d: _add_csv_table(
            d,
            "Tabla 1. Precisión media por condición (% aciertos)",
            bundle.rq1_group,
            [("Condición", "Condition"), ("Media", "MeanAccuracyPct"), ("N", "NParticipants")],
        ),
    )
    _with_full_width(
        doc,
        lambda d: _add_figure(
            d,
            bundle.figures_dir / "rq1_precision.png",
            "Figura 1. Precisión media por condición",
        ),
    )

    _with_full_width(
        doc,
        lambda d: _add_csv_table(
            d,
            "Tabla 2. Confianza media por condición (Unity, 1–7)",
            bundle.rq2_group,
            [("Condición", "Condition"), ("Media", "MeanConfidence"), ("N", "NParticipants")],
        ),
    )
    _with_full_width(
        doc,
        lambda d: _add_figure(
            d,
            bundle.figures_dir / "rq2_confidence.png",
            "Figura 2. Confianza media por condición",
        ),
    )

    _with_full_width(
        doc,
        lambda d: _add_csv_table(
            d,
            "Tabla 3. Brecha de calibración por condición",
            bundle.rq3_group,
            [("Condición", "Condition"), ("Brecha", "MeanCalibrationGap"), ("N", "NParticipants")],
        ),
    )
    _with_full_width(
        doc,
        lambda d: _add_figure(
            d,
            bundle.figures_dir / "rq3_calibration_gap.png",
            "Figura 3. Brecha de calibración",
        ),
    )

    if bundle.rq2_forms_tlx:
        _with_full_width(
            doc,
            lambda d: _add_csv_table(
                d,
                "Tabla 4. RAW-TLX por participante y bloque",
                bundle.rq2_forms_tlx,
                [
                    ("P##", "ParticipantCode"),
                    ("A", "RAW_TLX_A"),
                    ("B", "RAW_TLX_B"),
                    ("C", "RAW_TLX_C"),
                ],
            ),
        )

    if bundle.chat_group:
        _with_full_width(
            doc,
            lambda d: _add_csv_table(
                d,
                "Tabla 5. Métricas del agente (B vs C)",
                bundle.chat_group,
                [("Métrica", "Metric"), ("Media B", "Mean_B"), ("Media C", "Mean_C")],
            ),
        )
        for fig_name, caption in (
            ("chat_helpscore_b_vs_c.png", "Figura 4. HelpScore B vs C"),
            ("chat_exchanges_b_vs_c.png", "Figura 5. Intercambios de chat B vs C"),
            ("chat_leaks_b_vs_c.png", "Figura 6. Filtraciones del modelo B vs C"),
        ):
            _with_full_width(
                doc,
                lambda d, fn=fig_name, cap=caption: _add_figure(d, bundle.figures_dir / fn, cap),
            )

    if bundle.pilot_integrity:
        _with_full_width(
            doc,
            lambda d: _add_csv_table(
                d,
                "Tabla 6. Viabilidad técnica del estudio",
                bundle.pilot_integrity,
                [("Métrica", "Metric"), ("Valor", "Value")],
            ),
        )

    _add_section(doc, "4. Discusión", paper.sections.get("discusion", []))
    _add_section(doc, "5. Conclusiones", paper.sections.get("conclusiones", []))

    doc.add_heading("Referencias", level=1)
    for ref in paper.references:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ref)
        run.font.size = Pt(9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
