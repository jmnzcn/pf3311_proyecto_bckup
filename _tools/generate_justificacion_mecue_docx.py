#!/usr/bin/env python3
"""Genera DOCX de justificación del meCUE curado (PF-3311)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_SOURCE = ROOT / "docs" / "Justificacion_meCUE_Curado_PF3311_NeyFredJimenez_B03230.md"
OUT_PATH = ROOT / "docs" / "Justificacion_meCUE_Curado_PF3311_NeyFredJimenez_B03230.docx"

UCR_BLUE = RGBColor(0, 56, 101)
TEXT_DARK = RGBColor(33, 37, 41)
TEXT_MUTED = RGBColor(90, 98, 104)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK


def add_title(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = UCR_BLUE
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    p.space_after = Pt(6)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text).paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = cell
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    style_doc(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Justificación metodológica del meCUE curado")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = UCR_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Piloto PF-3311 · Efecto de la visibilidad de un agente virtual\n"
        "Ney Fred Jiménez Campos (B03230) · UCR\n\n"
        "Documento complementario al cuestionario meCUE en Google Forms"
    ).font.size = Pt(11)

    doc.add_page_break()

    add_title(doc, "1. Propósito", 2)
    add_body(
        doc,
        "Este documento justifica por qué y con qué criterio se aplicó una versión curada "
        "del cuestionario meCUE 2.0 (Minge & Thüring, 2018) en el piloto within-subjects "
        "(condiciones A, B, C), en lugar del instrumento completo sin adaptación.",
    )
    add_body(
        doc,
        "No sustituye la validación psicométrica del meCUE original; declara las decisiones "
        "de adaptación modular para maximizar validez de contenido y validez ecológica en "
        "un contexto experimental de sesión única.",
    )

    add_title(doc, "2. Marco: meCUE modular", 2)
    add_body(
        doc,
        "El meCUE 2.0 es un cuestionario modular (Minge, Thüring & Wagner, 2016; "
        "Minge & Thüring, 2018). Los módulos pueden seleccionarse según el objetivo del "
        "estudio, documentando qué ítems se administraron. Esto es coherente con la "
        "evaluación centrada en el usuario (ISO 9241-210:2019).",
    )

    add_title(doc, "3. Contexto del piloto", 2)
    add_body(
        doc,
        "«El producto» = el sistema de asistencia del bloque (chat en B; agente virtual en C). "
        "Administración inmediatamente después de cada bloque B o C.",
    )

    add_title(doc, "4. Criterios de curación", 2)
    add_bullet(doc, "Validez de contenido: pertinencia del ítem a las RQs (Haynes et al., 1995; Lawshe, 1975).")
    add_bullet(doc, "Validez ecológica: ítems respondibles sin simular mercado (Hornbæk, 2010; Lewis, 2021).")
    add_bullet(doc, "Diseño modular meCUE: subconjuntos documentados (Minge & Thüring, 2018).")
    add_bullet(doc, "Adaptación de instrumentos: ITC Guidelines (2017).")

    add_title(doc, "5. Resumen por módulo", 2)
    add_table(
        doc,
        ["Módulo", "Decisión", "Ítems"],
        [
            ["I", "Parcial (U2*, F2*)", "6"],
            ["II (solo C)", "Curado", "5"],
            ["III", "Completo", "12"],
            ["IV", "Curado", "2"],
            ["V", "Completo", "1"],
        ],
    )

    add_title(doc, "6. Módulo II — Curación (condición C)", 2)
    add_body(doc, "Omitidos S1–S3 (prestigio social): baja pertinencia en prototipo de laboratorio.")
    add_body(doc, "Conservados: A1–A3 (estética), C2 (cercanía). C3* adaptado: «Me costaría completar tareas similares sin un agente como este.»")

    add_title(doc, "6.4 Micro-adaptaciones Módulo I (B y C)", 2)
    add_table(
        doc,
        ["ID", "Decisión", "Enunciado administrado"],
        [
            ["U2*", "Adaptado", "Las funciones del asistente apoyan lo que necesitaba hacer en los casos de este bloque."],
            ["F2*", "Adaptado", "Con la ayuda de este asistente pude avanzar en los casos de este bloque."],
        ],
    )

    add_title(doc, "7. Módulo IV — Curación (B y C)", 2)
    add_table(
        doc,
        ["ID", "Decisión", "Enunciado / nota"],
        [
            ["IN.1", "Adaptado → IN.1*", "Volvería a usar un asistente como este para tareas similares."],
            ["IN.2", "Omitido", "Lenguaje de marketing; sesión única."],
            ["L1–L3", "Omitidos", "Lealtad comparativa; no hay alternativas reales."],
            ["IN.3", "Mantenido", "Al usar el producto, pierdo la noción del tiempo."],
        ],
    )
    add_body(
        doc,
        "En informes: «Módulo IV reducido (intención de reutilización + engagement)».",
    )

    add_title(doc, "8. Módulos I (parcial), III y V", 2)
    add_body(doc, "Módulo I: U2* y F2* adaptados; resto oficial. III y V sin cambios.")

    add_title(doc, "9. RAW-TLX", 2)
    add_body(doc, "Nota en formularios 1–3: no había límite de tiempo impuesto; TLX-T es presión subjetiva.")

    add_title(doc, "10. Limitaciones", 2)
    add_bullet(doc, "No se calculan puntuaciones normativas del meCUE completo ni lealtad oficial.")
    add_bullet(doc, "IN.1*, U2*, F2* y C3* son adaptaciones; analizar como ítems individuales.")
    add_bullet(doc, "Curación por juicio experimento; sin panel Lawshe formal en piloto.")
    add_bullet(doc, "meCUE complementa confianza por ítem en Unity (RQ2).")

    add_title(doc, "11. Redacción sugerida para informes", 2)
    add_body(
        doc,
        "Se administró un subconjunto modular curado del meCUE 2.0 (Minge & Thüring, 2018), "
        "documentado en este anexo, siguiendo criterios de validez de contenido (Haynes et al., 1995) "
        "y guías ITC (2017). Las puntuaciones son exploratorias, no baremos normativos.",
    )

    doc.add_page_break()
    add_title(doc, "Referencias bibliográficas", 2)
    refs = [
        "Brave, S., Nass, C., & Hutchinson, K. (2005). Computers that care. International Journal of Human-Computer Studies, 62(2), 161–178.",
        "Cassell, J., et al. (2000). Embodiment in conversational interfaces: Rea. CHI '00, 520–527.",
        "Hart, S. G. (2006). NASA-TLX; 20 years later. HFES Annual Meeting, 50(9), 904–908.",
        "Hart, S. G., & Staveland, L. E. (1988). Development of NASA-TLX. Advances in Psychology, 52, 139–183.",
        "Haynes, S. N., Richard, D. C. S., & Kubany, E. S. (1995). Content validity in psychological assessment. Psychological Assessment, 7(3), 238–247.",
        "Hornbæk, K. (2010). Dogmas in the assessment of usability evaluation methods. Behaviour & Information Technology, 29(1), 97–111.",
        "International Test Commission. (2017). ITC Guidelines for Translating and Adapting Tests (2nd ed.).",
        "ISO 9241-210:2019. Human-centred design for interactive systems.",
        "Lawshe, C. H. (1975). A quantitative approach to content validity. Personnel Psychology, 28(4), 563–575.",
        "Lewis, J. R. (2021). UX research and usability evaluation: The long view. Human–Computer Interaction, 36(5), 409–422.",
        "Minge, M., & Thüring, M. (2018). The MeCUE Questionnaire (2.0). Mensch und Computer 2018.",
        "Minge, M., Thüring, M., & Wagner, I. (2016). The meCUE questionnaire. i-com, 15(3), 237–248.",
    ]
    for ref in refs:
        add_body(doc, ref)

    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Generado: {OUT_PATH}")
    print(f"Fuente Markdown: {MD_SOURCE.name}")


if __name__ == "__main__":
    main()
