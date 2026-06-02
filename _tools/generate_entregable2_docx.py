#!/usr/bin/env python3
"""Genera Entregable 2 con diseño académico (Word + PDF vía Word COM)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "Entregable2_PF3311_NeyFredJimenez(B03230).md"
DOCX_PATH = ROOT / "docs" / "Entregable2_PF3311_NeyFredJimenez(B03230).docx"
PDF_PATH = ROOT / "docs" / "Entregable2_PF3311_NeyFredJimenez(B03230).pdf"

UCR_BLUE = RGBColor(0, 56, 101)
UCR_LIGHT = RGBColor(230, 240, 248)
TEXT_DARK = RGBColor(33, 37, 41)
TEXT_MUTED = RGBColor(90, 98, 104)
WHITE = RGBColor(255, 255, 255)
CODE_BG = RGBColor(248, 249, 250)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    add_page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for level, size, color in ((1, 16, UCR_BLUE), (2, 13, UCR_BLUE), (3, 12, TEXT_DARK)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_rich_text(paragraph, text: str, base_size=11, base_bold=False, base_italic=False) -> None:
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 40, 40)
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic
        run.font.size = Pt(base_size)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    add_rich_text(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    add_rich_text(p, text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    add_rich_text(p, text)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "003865")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    add_rich_text(p, text.strip("> ").strip('"'), base_italic=True)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8F9FA")
    set_cell_margins(cell, 100, 140, 100, 140)
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(33, 37, 41)
    doc.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) or c == "" for c in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    header = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "003865")
        set_cell_margins(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(re.sub(r"\*\*", "", text))
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    for i, row in enumerate(body, start=1):
        fill = "FFFFFF" if i % 2 == 1 else "E6F0F8"
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            set_cell_shading(cell, fill)
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_text(p, text, base_size=10)
    doc.add_paragraph()


def add_cover_page(doc: Document, meta: dict) -> None:
    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Universidad de Costa Rica")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = UCR_BLUE

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Programa de Posgrado en Computación e Informática")
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT_MUTED

    doc.add_paragraph()
    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = course.add_run("PF-3311 — Agentes Virtuales Inteligentes")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = TEXT_DARK

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("_" * 42)
    run.font.color.rgb = UCR_BLUE

    doc.add_paragraph()
    ent = doc.add_paragraph()
    ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ent.add_run("ENTREGABLE 2")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = UCR_BLUE

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Avance de Agente e Investigación")
    r.font.size = Pt(14)
    r.font.color.rgb = TEXT_DARK

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(meta.get("title", ""))
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = TEXT_DARK

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        ("Profesor:", meta.get("professor", "")),
        ("Investigador principal:", meta.get("author", "")),
        ("Fecha:", meta.get("date", "")),
        ("Ciclo:", meta.get("cycle", "")),
        ("Repositorio GitHub:", meta.get("github", "")),
        ("Video de demostración:", meta.get("video", "")),
    ]
    table = doc.add_table(rows=len(info_lines), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_lines):
        c0, c1 = table.rows[i].cells
        set_cell_margins(c0, 60, 80, 60, 80)
        set_cell_margins(c1, 60, 80, 60, 80)
        if i % 2 == 0:
            set_cell_shading(c0, "E6F0F8")
            set_cell_shading(c1, "E6F0F8")
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = UCR_BLUE
        add_rich_text(p1, value, base_size=10)

    doc.add_page_break()


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    add_rich_text(cap, caption, base_italic=True, base_size=10)
    for r in cap.runs:
        r.font.color.rgb = TEXT_MUTED


def build_from_markdown(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    meta = {
        "title": "Efecto de la visibilidad de un agente virtual en el desempeño y la confianza del usuario en tareas de decisión basadas en reglas",
        "professor": "Dr. Alexander Barquero",
        "author": "Ney Fred Jiménez Campos (B03230)",
        "date": "1 de junio de 2026",
        "cycle": "I Ciclo, 2026",
        "github": "Pendiente: URL en Mediación Virtual",
        "video": "Esc. 1 (A): https://youtu.be/ItcvsdxPfp8 | Esc. 2 (B): https://youtu.be/RY5pY_DVwDk | Esc. 3 (C): https://youtu.be/hLPTS9akSlg",
    }

    doc = Document()
    configure_document(doc)
    add_cover_page(doc, meta)

    skip_until_content = True
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        if skip_until_content:
            if line.startswith("## Resumen"):
                skip_until_content = False
            else:
                i += 1
                continue

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("|"):
            cells = parse_table_row(line)
            if is_separator_row(cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            if i >= len(lines) or not lines[i].startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            i += 1
            continue

        if line.startswith("## Tabla de contenidos"):
            doc.add_heading("Tabla de contenidos", level=1)
            p = doc.add_paragraph()
            add_toc(p)
            note = doc.add_paragraph()
            r = note.add_run("Actualizar campos en Word: clic derecho → Actualizar campo.")
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT_MUTED
            doc.add_page_break()
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if re.match(r"^\d+\.\s+\[", line.strip()):
            m = re.match(r"^\d+\.\s+\[(.+?)\]", line.strip())
            if m:
                add_numbered(doc, m.group(1))
            i += 1
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            i += 1
            continue

        if line.startswith("> "):
            add_quote(doc, line)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        if line.startswith("*") and line.strip().endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, line.strip(), base_italic=True, base_size=10)
            i += 1
            continue

        if line.startswith("**Figura") or line.startswith("**Tabla"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            add_rich_text(p, line.strip(), base_italic=True, base_size=10)
            i += 1
            continue

        if line.strip().startswith("*(") and line.strip().endswith(")*"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, line.strip()[1:-1], base_italic=True, base_size=9)
            for run in p.runs:
                run.font.color.rgb = TEXT_MUTED
            i += 1
            continue

        img_match = re.match(r"^!\[(.+?)\]\((.+?)\)\s*$", line.strip())
        if img_match:
            caption, rel = img_match.groups()
            img_path = (md_path.parent / rel).resolve()
            if img_path.is_file():
                add_image(doc, img_path, caption)
            else:
                add_body(doc, f"[Imagen no encontrada: {rel}]")
            i += 1
            continue

        add_body(doc, line)
        i += 1

    doc.save(docx_path)
    print(f"DOCX: {docx_path}")


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    try:
        import comtypes.client  # type: ignore

        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"PDF: {pdf_path}")
    except Exception as exc:
        print(f"PDF no generado automáticamente: {exc}")
        print("Abrí el DOCX en Word → Guardar como PDF.")


def main() -> int:
    md = MD_PATH
    docx = DOCX_PATH
    pdf = PDF_PATH
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        docx = Path(sys.argv[2])
    figures_dir = md.parent / "figures"
    if not figures_dir.is_dir() or not any(figures_dir.glob("figura_*.png")):
        print("Figuras no encontradas; ejecutá: python _tools/generate_entregable_figures.py")
    build_from_markdown(md, docx)
    export_pdf(docx.resolve(), pdf.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
