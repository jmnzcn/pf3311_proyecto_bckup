#!/usr/bin/env python3
"""Genera cuestionarios alineados con E1, E2, Tarea 10 y protocolo PF-3311."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "Cuestionario_meCUE_PF3311_NeyFredJimenez_B03230.docx"
UCR_COPY = Path(
    r"c:\UCR Maestria\2026\Semestre 1\PF-3311 Temas Especiales de Ingeniería de Sistemas de Información Agentes Virtuales Inteligentes\Proyecto Investigacion\Preguntas Proyecto Investigacion\Cuestionario_meCUE_PF3311_NeyFredJimenez_B03230.docx"
)

UCR_BLUE = RGBColor(0, 56, 101)
TEXT_DARK = RGBColor(33, 37, 41)
TEXT_MUTED = RGBColor(90, 98, 104)

MODULE_I = [
    ("U1", "El producto es fácil de usar."),
    ("U2", "Las funciones del asistente apoyan lo que necesitaba hacer en los casos de este bloque."),
    ("U3", "Es evidente rápidamente cómo usar el producto."),
    ("F1", "Considero que el producto es extremadamente útil."),
    ("U4", "Los procedimientos de uso del producto son sencillos de entender."),
    ("F2", "Con la ayuda de este asistente pude avanzar en los casos de este bloque."),
]

# Curado para laboratorio: estética + commitment; sin ítems de prestigio social (S1–S3).
MODULE_II_CURATED = [
    ("A1", "El producto está diseñado de forma creativa."),
    ("A2", "El diseño se ve atractivo."),
    ("A3", "El producto es elegante / tiene estilo."),
    ("C2", "El personaje del agente me resultó cercano."),
    ("C3", "Me costaría completar tareas similares sin un agente como este."),
]

MODULE_III = [
    ("PA.1", "El producto me entusiasma."),
    ("ND.1", "El producto me cansa."),
    ("NA.1", "El producto me molesta."),
    ("PD.1", "El producto me relaja."),
    ("ND.2", "Al usar este producto me siento agotado/a."),
    ("PD.2", "El producto me hace sentir feliz."),
    ("NA.2", "El producto me frustra."),
    ("PA.2", "El producto me hace sentir eufórico/a."),
    ("ND.3", "El producto me hace sentir pasivo/a."),
    ("PD.3", "El producto me calma."),
    ("PA.3", "Al usar este producto me siento alegre."),
    ("NA.3", "El producto me enoja."),
]

# Curado para laboratorio: intención de reutilización + engagement (ver Justificacion_meCUE_Curado_…docx).
MODULE_IV_CURATED = [
    ("IN.1*", "Volvería a usar un asistente como este para tareas similares."),
    ("IN.3", "Al usar el producto, pierdo la noción del tiempo."),
]

# Ítems oficiales omitidos en Módulo IV: IN.2, L1, L2, L3 (lealtad / marketing; baja validez ecológica).

# RAW-TLX inspirado en Hart & Staveland (1988); escala 1–7 en el piloto.
NASA_RAW_TLX = [
    ("TLX-M", "La tarea me exigió mucha actividad mental y concentración."),
    ("TLX-T", "Sentí presión de tiempo mientras resolvía los casos."),
    ("TLX-E", "Tuve que trabajar muy duro (esfuerzo) para completar la tarea."),
    ("TLX-F", "Me sentí frustrado/a, tenso/a o irritado/a durante el bloque."),
    ("TLX-P", "Me sentí seguro/a de cómo desempeñé la tarea en este bloque."),
    ("TLX-D", "La tarea me resultó exigente en general."),
]

AGE_RANGES = [
    "18–24 años",
    "25–34 años",
    "35–44 años",
    "45–54 años",
    "55–64 años",
    "65 años o más",
]

PROFILE_ITEMS = [
    ("P1", "¿En qué rango de edad te encontrás?"),
    ("P2", "Nivel educativo más alto alcanzado:"),
    ("P3", "¿Con qué frecuencia usa asistentes digitales (Siri, Alexa, ChatGPT, etc.)?"),
    ("P4", "¿Ha interactuado antes con agentes virtuales o avatares conversacionales?"),
]

MECUE_INSTRUCTIONS = (
    "Respondé según tu experiencia con el sistema indicado en este bloque. "
    "No hay respuestas correctas. Tus respuestas son anónimas (código P01, P02, etc.). "
    "Escala 1 (totalmente en desacuerdo) a 7 (totalmente de acuerdo). Decidí de forma espontánea."
)

PRODUCT_NOTE = (
    "«El producto» = el sistema de asistencia del bloque que acabás de completar."
)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK


def add_title(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = UCR_BLUE
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    p.space_after = Pt(6)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = TEXT_MUTED
    p.runs[0].font.size = Pt(10)
    p.space_after = Pt(10)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text).paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_likert_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(
        ["ID", "Ítem", "1", "2", "3", "4", "5", "6", "7", "Respuesta"]
    ):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for item_id, statement in items:
        row = table.add_row().cells
        row[0].text = item_id
        row[1].text = statement
        for c in range(2, 9):
            row[c].text = "○"
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_module_v_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=13)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "V1"
    table.rows[0].cells[1].text = (
        "¿Cómo evaluás el producto (el sistema de asistencia) en general?"
    )
    for i, label in enumerate(
        ["−5", "−4", "−3", "−2", "−1", "0", "+1", "+2", "+3", "+4", "+5"],
        start=2,
    ):
        table.rows[0].cells[i].text = label
        table.rows[1].cells[i].text = "○"
    doc.add_paragraph()


def add_meta_fields(doc: Document, condition: str) -> None:
    add_body(
        doc,
        f"Condición: {condition}     Código participante (P01–P15): _______________",
    )
    doc.add_paragraph()


def add_consistency_matrix(doc: Document) -> None:
    add_title(doc, "Matriz de consistencia (Entregable 2 / Tarea 10)", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["RQ", "Constructo", "Instrumento (este documento)", "Momento"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("RQ1", "Precisión", "CSV Unity (AnswerLetter vs CorrectAnswerLetter)", "Cada ítem"),
        ("RQ2", "Confianza por decisión", "App Unity (estrellas 1–7)", "Tras cada respuesta"),
        ("RQ2", "UX / emociones / evaluación global", "meCUE 2.0 modular (Forms 2 y 3)", "Tras bloques B y C"),
        ("RQ3", "Calibración", "Derivada: confianza normalizada − acierto", "Análisis post hoc CSV"),
        ("Apoyo", "Carga cognitiva", "RAW-TLX (Forms 1–3)", "Tras cada bloque A/B/C"),
        ("Control", "Perfil", "Formulario 0", "Inicio de sesión"),
    ]
    for rq, construct, instrument, moment in rows:
        row = table.add_row().cells
        row[0].text = rq
        row[1].text = construct
        row[2].text = instrument
        row[3].text = moment
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    style_doc(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Instrumentos de evaluación — Piloto PF-3311")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = UCR_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Alineado con Entregable 1 (RQs), Entregable 2 (métricas/instrumentos),\n"
        "Tarea 10 (matriz metodológica) y Protocolo del investigador\n\n"
        "meCUE 2.0 modular curado (Minge & Thüring, 2018) · RAW-TLX · Perfil\n"
        "Ver: Justificacion_meCUE_Curado_PF3311_NeyFredJimenez_B03230.docx\n"
        "Investigador: Ney Fred Jiménez Campos (B03230) · UCR · PF-3311"
    ).font.size = Pt(11)

    doc.add_page_break()

    add_title(doc, "Guía para el investigador", 2)
    add_body(
        doc,
        "Administración según el diseño within-subjects (A, B, C) y la matriz de Tarea 10. "
        "Trasladar cada formulario a Google Forms antes del piloto.",
    )
    add_bullet(doc, "Formulario 0 — Inicio: cuestionario de perfil + mismo código P01 ingresado en Unity.")
    add_bullet(doc, "Formulario 1 — Tras bloque A: RAW-TLX (sin meCUE; no hay sistema de asistencia).")
    add_bullet(doc, "Formulario 2 — Tras bloque B: meCUE I + III + IV curado (2 ítems) + V + RAW-TLX.")
    add_bullet(doc, "Formulario 3 — Tras bloque C: meCUE I + II curado + III + IV curado + V + RAW-TLX.")
    add_bullet(doc, "Justificación de la curación: docs/Justificacion_meCUE_Curado_PF3311_NeyFredJimenez_B03230.docx")
    add_bullet(doc, "Confianza por ítem (1–7): solo en Unity; no repetir en estos formularios.")
    add_bullet(doc, "Documento fuente: docs/Cuestionario_meCUE_PF3311_NeyFredJimenez_B03230.docx")
    add_body(
        doc,
        "Citación: «Se aplicó un subconjunto modular curado de meCUE 2.0 (Minge & Thüring, 2018), "
        "documentado en Justificacion_meCUE_Curado_PF3311_NeyFredJimenez_B03230, traducido al español, "
        "inmediatamente después de cada bloque B y C, conforme a la matriz de consistencia del Entregable 2.»",
    )
    add_consistency_matrix(doc)
    doc.add_page_break()

    # Form 0
    add_title(doc, "Formulario 0 — Cuestionario de perfil (inicio de sesión)", 1)
    add_subtitle(doc, "Protocolo §2.3 · ~3 min · Google Forms separado")
    add_meta_fields(doc, "—")
    add_body(doc, "Instrucciones: Respondé con sinceridad. Mismo código P01–P15 que usarás en la aplicación.")
    for pid, label in PROFILE_ITEMS:
        if pid == "P1":
            add_body(doc, f"{pid}. {label}")
            for option in AGE_RANGES:
                add_bullet(doc, option)
        else:
            add_body(doc, f"{pid}. {label} ________________________________")
    doc.add_page_break()

    # Form A
    add_title(doc, "Formulario 1 — Tras condición A", 1)
    add_subtitle(doc, "~4 min")
    add_meta_fields(doc, "A")
    add_body(doc, "Carga cognitiva de la tarea (sin evaluar sistema de asistencia). Escala 1–7.")
    add_body(
        doc,
        "Nota: no había límite de tiempo impuesto en la aplicación; respondé según cómo te sentiste durante el bloque.",
    )
    add_title(doc, "RAW-TLX adaptado (6 ítems)", 2)
    add_likert_table(doc, NASA_RAW_TLX)
    doc.add_page_break()

    # Form B
    add_title(doc, "Formulario 2 — Tras condición B (chat)", 1)
    add_subtitle(doc, "~10 min")
    add_meta_fields(doc, "B")
    add_body(doc, MECUE_INSTRUCTIONS)
    add_body(doc, PRODUCT_NOTE + " (asistencia por chat de texto).")
    add_body(
        doc,
        "Nota RAW-TLX: no había límite de tiempo impuesto; en carga cognitiva, respondé según cómo te sentiste.",
    )
    add_title(doc, "Módulo I — Cualidades instrumentales (6 ítems)", 2)
    add_likert_table(doc, MODULE_I)
    add_title(doc, "Módulo III — Emociones (12 ítems)", 2)
    add_likert_table(doc, MODULE_III)
    add_title(doc, "Módulo IV — Consecuencias de uso (2 ítems curados)", 2)
    add_body(doc, "Ver justificación: omitidos L1–L3 e IN.2; IN.1 adaptado a reutilización de asistente similar.")
    add_likert_table(doc, MODULE_IV_CURATED)
    add_title(doc, "Módulo V — Evaluación global (1 ítem, escala −5 a +5)", 2)
    add_module_v_table(doc)
    add_title(doc, "RAW-TLX adaptado (6 ítems)", 2)
    add_likert_table(doc, NASA_RAW_TLX)
    doc.add_page_break()

    # Form C
    add_title(doc, "Formulario 3 — Tras condición C (agente virtual)", 1)
    add_subtitle(doc, "~12 min")
    add_meta_fields(doc, "C")
    add_body(doc, MECUE_INSTRUCTIONS)
    add_body(doc, PRODUCT_NOTE + " (agente virtual con chat y voz).")
    add_body(
        doc,
        "Nota RAW-TLX: no había límite de tiempo impuesto; en carga cognitiva, respondé según cómo te sentiste.",
    )
    add_title(doc, "Módulo I — Cualidades instrumentales (6 ítems)", 2)
    add_likert_table(doc, MODULE_I)
    add_title(doc, "Módulo II — Cualidades no instrumentales (5 ítems curados)", 2)
    add_body(doc, "Incluye estética y cercanía del avatar (embodiment). Sin ítems de prestigio social.")
    add_likert_table(doc, MODULE_II_CURATED)
    add_title(doc, "Módulo III — Emociones (12 ítems)", 2)
    add_likert_table(doc, MODULE_III)
    add_title(doc, "Módulo IV — Consecuencias de uso (2 ítems curados)", 2)
    add_body(doc, "Ver justificación: omitidos L1–L3 e IN.2; IN.1 adaptado a reutilización de asistente similar.")
    add_likert_table(doc, MODULE_IV_CURATED)
    add_title(doc, "Módulo V — Evaluación global (1 ítem)", 2)
    add_module_v_table(doc)
    add_title(doc, "RAW-TLX adaptado (6 ítems)", 2)
    add_likert_table(doc, NASA_RAW_TLX)

    doc.add_page_break()
    add_title(doc, "Referencias", 2)
    add_body(
        doc,
        "Minge, M., & Thüring, M. (2018). The MeCUE Questionnaire (2.0). Springer. https://www.mecue.de",
    )
    add_body(
        doc,
        "Hart, S. G., & Staveland, L. E. (1988). NASA-TLX. Advances in Psychology, 52, 139–183.",
    )
    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Generado: {OUT_PATH}")
    if UCR_COPY.parent.exists():
        doc.save(UCR_COPY)
        print(f"Copiado: {UCR_COPY}")


if __name__ == "__main__":
    main()
