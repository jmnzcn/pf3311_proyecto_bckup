#!/usr/bin/env python3
"""Genera docs/MAPA_CODIGO.docx — descripción de cada archivo de código del proyecto."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "MAPA_CODIGO.docx"

UCR_BLUE = RGBColor(0, 56, 101)
TEXT_DARK = RGBColor(33, 37, 41)
TEXT_MUTED = RGBColor(90, 98, 104)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for level, size in ((1, 16), (2, 13), (3, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = UCR_BLUE if level <= 2 else TEXT_DARK
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)


def add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(1.0 * level)
    p.add_run(text)


def add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(label + " ")
    r1.bold = True
    p.add_run(value)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8F9FA")
    set_cell_margins(cell, 100, 140, 100, 140)
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "003865")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            set_cell_margins(cell, 60, 100, 60, 100)
            cell.paragraphs[0].add_run(text).font.size = Pt(10)

    doc.add_paragraph()


def add_casual_paragraph(doc: Document, text: str) -> None:
    """Párrafo explicativo en tono conversacional (sin negritas innecesarias)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_interaction_guide(doc: Document) -> None:
    """Guía larga: chat, voz y avatar explicados para humanos no técnicos."""
    doc.add_heading("Cómo funciona el chat, la voz y el avatar (explicación detallada)", level=1)

    add_casual_paragraph(
        doc,
        "Esta sección cuenta el recorrido completo como si lo vieras desde la silla del participante, "
        "y después dice qué pieza de código hace cada paso. La idea es que entiendas el flujo sin tener "
        "que leer Unity línea por línea.",
    )

    doc.add_heading("Las tres condiciones en una frase", level=2)
    add_casual_paragraph(
        doc,
        "El experimento tiene tres “modos” que el participante elige al inicio (botones INICIAR). "
        "No son tres programas distintos: es el mismo juego con partes encendidas o apagadas.",
    )
    add_bullet(doc, "Condición A (escenario 0): solo preguntas. No hay chat, no hay personaje, no hay voz.")
    add_bullet(doc, "Condición B (escenario 1): aparece el chat. El agente responde por texto (Google Gemini).")
    add_bullet(
        doc,
        "Condición C (escenario 2): chat + personaje 3D en el panel derecho + el mismo texto del agente "
        "se convierte en voz (Azure) y la boca del avatar se mueve.",
    )

    doc.add_heading("Qué ve el participante antes de chatear", level=2)
    add_casual_paragraph(
        doc,
        "Primero está la pantalla de consentimiento (casilla + Continuar). Ahí ExperimentLogic genera "
        "un ID de sesión único, por ejemplo ID-20260602143000-1234, que después va en todos los CSV.",
    )
    add_casual_paragraph(
        doc,
        "Luego elige A, B o C. QuestionManager.BeginScenario() carga las 6 preguntas de ese escenario "
        "y decide qué UI mostrar: en A oculta botón de ayuda, caja de texto y panel de chat; en B y C los muestra; "
        "solo en C activa el GameObject del avatar (TTBoyB).",
    )

    doc.add_heading("Flujo del chat (condiciones B y C)", level=2)
    add_casual_paragraph(
        doc,
        "Mientras responde una pregunta, el participante puede escribir en el cuadro de texto y pulsar "
        "el botón de pedir ayuda (en código: AskForHelp en ExperimentLogic). Eso dispara todo lo del agente.",
    )

    add_body(doc, "Paso a paso — qué pasa cuando envía un mensaje:", bold=True)
    steps_chat = [
        (
            "1. Validaciones rápidas",
            "Si está en condición A, AskForHelp no hace nada. Si ya hay una petición en curso (geminiInFlight), "
            "tampoco deja mandar otro mensaje hasta que termine (evita doble clic y respuestas mezcladas). "
            "Si el cuadro está vacío, ignora el envío.",
        ),
        (
            "2. Mensaje “Enviando…” en pantalla",
            "Antes de llamar a internet, el chat muestra el texto del estudiante con una línea Enviando…. "
            "Eso es RefreshPendingStudentMessage. Si la API falla, ese mensaje se borra y vuelve el historial "
            "anterior (RevertPendingStudentMessage), para que el participante no crea que el mensaje “quedó guardado”.",
        ),
        (
            "3. Armar el “briefing” del profesor (system prompt)",
            "ExperimentLogic arma un texto largo con BuildSystemPreamble(). Le dice al modelo: "
            "sos un profesor empático; acá está el enunciado de la pregunta actual; NUNCA des la letra correcta; "
            "solo guía con pistas cortas; no saludes si el estudiante no saludó. "
            "El enunciado sale de question.situation (la misma situación que se ve arriba en la pantalla).",
        ),
        (
            "4. Llamada a Gemini",
            "Se hace un POST a gemini-2.5-flash con la clave apiKey. El historial de la pregunta actual va en apiTurns "
            "(máximo 30 turnos). La primera vez el prompt va pegado al mensaje del estudiante; en mensajes siguientes "
            "se manda la conversación como turnos user/model.",
        ),
        (
            "5. Si todo sale bien",
            "Se confirma el mensaje del estudiante en el historial visible, llega el texto del “Profesor”, "
            "se guarda en apiTurns y se ejecuta ProcessResponse().",
        ),
        (
            "6. Si algo falla",
            "Respuesta vacía, error de red, demasiadas peticiones (429) o servidor caído (503 con reintentos): "
            "mensaje amarillo en el chat y el mensaje del estudiante no entra al contexto de Gemini, "
            "así puede reintentar sin que el modelo “recuerde” un envío que en realidad no llegó.",
        ),
    ]
    for title, detail in steps_chat:
        p = doc.add_paragraph()
        r1 = p.add_run(title + " — ")
        r1.bold = True
        p.add_run(detail)

    doc.add_heading("Qué “personalidad” tiene el agente en el chat", level=2)
    add_casual_paragraph(
        doc,
        "No es un personaje con memoria infinita: es Gemini con reglas escritas a mano en BuildSystemPreamble(). "
        "En la práctica se comporta como un profesor que quiere que pienses, no como un buscador que te da la "
        "respuesta del examen. Si el participante pide “dame la B”, el prompt le obliga a redirigir con una "
        "pregunta guía sobre el caso.",
    )
    add_casual_paragraph(
        doc,
        "Cada vez que pasás a la siguiente pregunta, QuestionManager llama ClearChatHistory(): el chat visible "
        "vuelve al placeholder gris y apiTurns se vacía. O sea: la conversación no cruza de una pregunta a otra; "
        "cada ítem empieza de cero con el agente.",
    )

    doc.add_heading("Flujo de la voz y el avatar (solo condición C)", level=2)
    add_casual_paragraph(
        doc,
        "En B solo leés al agente. En C pasa lo mismo en el chat, pero además “habla” el muñeco del panel derecho. "
        "Importante: la voz no es otra IA distinta. Es el mismo texto que Gemini ya escribió, leído en voz alta por Azure.",
    )

    add_body(doc, "Cadena completa texto → voz → boca:", bold=True)
    add_code_block(
        doc,
        [
            "Participante escribe → Gemini responde texto",
            "         ↓",
            "ProcessResponse() muestra texto en el chat",
            "         ↓",
            "¿Avatar activo? (characterModel.activeInHierarchy)",
            "    Sí → AgentSpeechController.Speak(texto)",
            "         ↓",
            "    AzureLipSync.SpeakText(texto)",
            "         ↓",
            "    Azure devuelve audio PCM + eventos de visemas",
            "         ↓",
            "    AudioSource reproduce + blend shapes mueven la boca",
            "         ↓",
            "    Animator alterna idle / hablando (TTB_idle1, TTB_talk2)",
        ],
    )

    voice_steps = [
        (
            "ProcessResponse y cuándo habla",
            "Después de limpiar markdown del texto (*, #, comillas), lo muestra como Profesor: en el panel. "
            "Luego busca AgentSpeechController en GameManager. Solo llama Speak() si el avatar está activo "
            "(escenario C). En B nunca se llama Speak aunque el componente exista en la escena.",
        ),
        (
            "AgentSpeechController (intermediario chico)",
            "Es un script de unas líneas: recibe el string y se lo pasa a AzureLipSync. Existe para no "
            "mezclar lógica de chat (ExperimentLogic) con la de síntesis de voz.",
        ),
        (
            "AzureLipSync — la voz",
            "Usa Microsoft Cognitive Services con voz es-MX-JorgeNeural y región eastus. "
            "SpeakTextAsync genera audio crudo 16 kHz; Unity lo convierte a AudioClip y lo reproduce "
            "por el AudioSource del personaje.",
        ),
        (
            "AzureLipSync — la boca",
            "Mientras sintetiza, Azure manda eventos VisemeReceived (qué forma tiene la boca en cada instante). "
            "El script los guarda en una lista y en LateUpdate mueve blend shapes A, E, U y mandíbula del mesh "
            "del personaje para que parezca que articula en español.",
        ),
        (
            "Animación del cuerpo",
            "Además de la boca, el Animator recibe isTalking y hace crossfade entre idle y talk. "
            "Cuando termina el audio (por tiempo del clip, no solo por isPlaying), vuelve a idle y resetea la boca.",
        ),
        (
            "Si Azure falla",
            "No rompe el experimento: avisa en el chat con AppendChatNotice (podés seguir leyendo la respuesta). "
            "El texto de Gemini ya quedó en pantalla.",
        ),
        (
            "Pausa si cambiás de ventana",
            "Si el participante alt+tab o minimiza, OnApplicationFocus pausa el audio y al volver lo reanuda "
            "si no terminó. Está pensado para la VM del piloto, no para cortar la frase a la mitad sin querer.",
        ),
    ]
    for title, detail in voice_steps:
        p = doc.add_paragraph()
        r1 = p.add_run(title + " — ")
        r1.bold = True
        p.add_run(detail)

    doc.add_heading("Cómo se ve el avatar en el panel (AvatarDisplayController)", level=2)
    add_casual_paragraph(
        doc,
        "El modelo 3D (TTBoyB) vive en la escena, pero en condición C no lo ves “flotando” en el mundo del juego "
        "como un NPC. AvatarDisplayController crea una cámara chica y un RenderTexture: renderiza solo al personaje "
        "en una capa especial (AvatarLayer) y esa imagen la pega en el RawImage del panel derecho de la UI.",
    )
    add_casual_paragraph(
        doc,
        "Cuando QuestionManager enciende characterModel, LateUpdate detecta el cambio y encuadra al personaje "
        "mirando hacia la cámara principal. Si cerrás C o terminás el escenario, el modelo se apaga y el panel "
        "deja de mostrarlo.",
    )

    doc.add_heading("Diagrama general (todo junto)", level=2)
    add_code_block(
        doc,
        [
            "[Inicio] Consentimiento → ID sesión",
            "     ↓",
            "Elegir A / B / C  →  QuestionManager.BeginScenario",
            "     ↓",
            "Loop por cada pregunta:",
            "   · Leer enunciado · elegir A-D · confianza · Entregar → DataLogger (CSV principal)",
            "   · Si B o C: opcionalmente chat → Gemini → texto en pantalla",
            "   · Si C además: mismo texto → Azure TTS + lip-sync + panel avatar",
            "     ↓",
            "Siguiente pregunta (chat se limpia) o pantalla final / Otro escenario",
        ],
    )

    doc.add_heading("Quién llama a quién (referencia rápida)", level=2)
    add_table(
        doc,
        ["Acción del participante", "Script que reacciona", "Qué hace"],
        [
            ["Marca consentimiento y Continuar", "ExperimentLogic", "Crea UserID, muestra selección de escenario"],
            ["Pulsa INICIAR escenario B o C", "QuestionManager + ExperimentLogic", "Carga preguntas; muestra chat; en C activa avatar"],
            ["Escribe y pide ayuda", "ExperimentLogic.AskForHelp", "Coroutine CallGemini → API Google"],
            ["Recibe respuesta del agente", "ExperimentLogic.ProcessResponse", "Actualiza chat; en C dispara Speak"],
            ["Escucha al personaje", "AgentSpeechController → AzureLipSync", "TTS + visemas + animación"],
            ["Ve al personaje en el panel", "AvatarDisplayController", "RenderTexture del modelo 3D"],
            ["Entrega respuesta con estrellas", "QuestionManager → DataLogger", "CSV en CSV data/ (dato principal)"],
            ["Pasa a otra pregunta", "QuestionManager.ShowNextQuestion", "ClearChatHistory si hay chat"],
        ],
    )

    doc.add_heading("Datos que quedan guardados del chat", level=2)
    add_casual_paragraph(
        doc,
        "Para el análisis estadístico principal usás CSV data/ (DataLogger): respuesta elegida, si acertó, "
        "confianza, tiempo por pregunta. El chat completo no va ahí.",
    )
    add_casual_paragraph(
        doc,
        "ExperimentLogic además puede escribir en Logs/ExperimentData_{ID}.csv acciones como mensajes del chat "
        "(exportación secundaria / legado). Si necesitás auditar conversaciones, ese archivo o logs de sesión "
        "son el rastro; el diseño del estudio se centra en desempeño y confianza en el CSV principal.",
    )

    doc.add_heading("Dónde está el prompt del agente (para editarlo)", level=2)
    add_casual_paragraph(
        doc,
        "El “prompt” no es un archivo .txt aparte: está escrito en código C# y se arma cada vez que el "
        "participante envía un mensaje al chat.",
    )
    add_label_value(doc, "Archivo:", "Assets/_Project/Scripts/Core/ExperimentLogic.cs")
    add_label_value(doc, "Método:", "BuildSystemPreamble(string scenarioText) — aprox. líneas 190–198")
    add_label_value(
        doc,
        "Cuándo se usa:",
        "CallGemini() lo llama y lo manda a Gemini junto con el mensaje del estudiante "
        "(vía BuildGeminiRequestJson / BuildConversationContents).",
    )
    add_body(doc, "Qué contiene ese texto (resumen):", bold=True)
    add_bullet(doc, "Prefijo: «Contexto del Escenario Actual:» + el enunciado de la pregunta activa.")
    add_bullet(
        doc,
        "Reglas fijas: actuar como profesor empático; nunca dar la opción correcta; guiar con pistas; "
        "no saludar si el estudiante no saludó; respuestas cortas (2–3 oraciones).",
    )
    add_casual_paragraph(
        doc,
        "El enunciado que se inyecta en el prompt NO está en ExperimentLogic: sale de "
        "QuestionManager → lista scenarios → ExperimentQuestion.situation (se edita en el Inspector de Unity "
        "o se regenera con _tools/generate_scenarios_yaml.py).",
    )
    add_casual_paragraph(
        doc,
        "Para cambiar cómo se comporta el agente (tono, prohibiciones, longitud), editá las cadenas de texto "
        "dentro de BuildSystemPreamble(). Para cambiar el caso que lee, editá la pregunta en QuestionManager.",
    )

    doc.add_heading("Dónde están las API keys (para ejecutar el proyecto)", level=2)
    add_casual_paragraph(
        doc,
        "Hay dos servicios externos: Google Gemini (chat en B y C) y Azure Speech (voz en C). "
        "En el repositorio de GitHub los campos suelen estar en placeholder (YOUR_*); para correr en Unity "
        "o en un build local tenés que poner claves válidas.",
    )
    add_table(
        doc,
        ["Servicio", "GameObject en la escena", "Componente", "Campos en el Inspector"],
        [
            ["Google Gemini", "GameManager", "ExperimentLogic", "apiKey, endpoint"],
            ["Azure Speech (TTS + visemas)", "TTBoyB", "AzureLipSync", "subscriptionKey, region (ej. eastus)"],
        ],
    )
    add_label_value(
        doc,
        "Escena donde se guardan los valores al correr en Editor:",
        "Assets/Scenes/SampleScene.unity (Unity serializa apiKey y subscriptionKey ahí).",
    )
    add_label_value(
        doc,
        "Valores por defecto en código (si la escena no los sobreescribe):",
        "ExperimentLogic.cs → apiKey = \"YOUR_GEMINI_API_KEY_HERE\"; "
        "AzureLipSync.cs → subscriptionKey = \"YOUR_API_KEY_HERE\".",
    )
    add_casual_paragraph(
        doc,
        "Forma práctica de configurar: abrís SampleScene.unity, seleccionás GameManager o TTBoyB en la "
        "jerarquía, y en el Inspector completás los campos del componente. Eso es lo que usa Play y el .exe "
        "al compilar (las claves quedan embebidas en el build si no las limpiás antes).",
    )
    add_body(doc, "Antes de subir a GitHub o entregar código fuente:", bold=True)
    add_bullet(
        doc,
        "Ejecutar python _tools/clean_for_delivery.py — reemplaza claves reales por YOUR_GEMINI_API_KEY_HERE "
        "y YOUR_API_KEY_HERE en ExperimentLogic.cs y SampleScene.unity.",
    )
    add_bullet(doc, "No commitear claves en el repo (requisito del curso). El zip Build/Windows/ puede llevar claves solo para el evaluador.")
    add_casual_paragraph(
        doc,
        "AgentSpeechController no tiene API keys: solo reenvía texto a AzureLipSync.",
    )

    doc.add_paragraph()


def add_file_section(doc: Document, path: str, role: str, bullets: list[str], extra: str | None = None) -> None:
    doc.add_heading(path, level=3)
    add_label_value(doc, "Rol:", role)
    add_body(doc, "Responsabilidades:", bold=True)
    for b in bullets:
        add_bullet(doc, b)
    if extra:
        add_body(doc, extra)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_heading("Mapa del código — ExperimentPrototypeB03230", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(
        doc,
        "Documento de referencia que describe qué hace cada archivo de código del proyecto Unity "
        "(PF-3311, investigador Ney Fred Jiménez Campos, B03230).",
    )
    add_body(doc, "Escena principal: Assets/Scenes/SampleScene.unity")
    add_body(doc, "Unity: 6000.3.11f1 · Plataforma: Windows standalone (pantalla completa 1920×1080)")

    # --- Resumen ---
    doc.add_heading("Resumen del flujo", level=1)
    add_body(
        doc,
        "QuestionManager controla preguntas y UI. ExperimentLogic gestiona consentimiento y chat con Gemini. "
        "DataLogger guarda respuestas en CSV data/. En condiciones B y C el participante puede pedir ayuda al agente; "
        "en C además AzureLipSync sintetiza voz y AvatarDisplayController muestra el personaje 3D.",
    )

    add_table(
        doc,
        ["Condición", "Escenario", "Chat Gemini", "Voz Azure", "Avatar 3D"],
        [
            ["A", "Índice 0", "No", "No", "No"],
            ["B", "Índice 1", "Sí (solo texto)", "No", "No"],
            ["C", "Índice 2", "Sí", "Sí", "Sí (panel derecho)"],
        ],
    )

    add_interaction_guide(doc)

    # --- Scripts ---
    doc.add_heading("Scripts del experimento (Assets/_Project/)", level=1)
    add_body(doc, "Código escrito para este proyecto. Se adjuntan como componentes en SampleScene.unity.")

    add_file_section(
        doc,
        "Scripts/Core/QuestionManager.cs",
        "Motor principal del flujo experimental y de la interfaz de preguntas.",
        [
            "Define escenarios (ScenarioDefinition) y preguntas (ExperimentQuestion) con enunciado, opciones A–D y respuesta correcta.",
            "Flujo: elegir escenario → pregunta → opción → confianza (estrellas 1–5) → entregar → siguiente o pantalla final.",
            "Escenario A: sin chat ni avatar. B: chat visible. C: chat + modelo 3D.",
            "Barra de progreso, contador de preguntas, enlace a encuesta y popups de salida segura.",
            "Al entregar, llama a DataLogger.SaveAnswer() (CSV principal) y ExperimentLogic.SaveDataToCSV() (Logs/).",
            "SafeExit(): cierra la app y borra CSV parcial si la sesión quedó incompleta en el primer escenario.",
        ],
        extra="Namespace: MyProject",
    )

    add_file_section(
        doc,
        "Scripts/Core/ExperimentLogic.cs",
        "Consentimiento informado, ID de sesión y agente conversacional (Google Gemini).",
        [
            "Pantalla inicial: consentimiento, botón continuar, ID de sesión (ID-yyyyMMddHHmmss-xxxx).",
            "OnScenarioSelected(int): inicia escenario 0, 1 o 2 desde los botones INICIAR.",
            "AskForHelp() → CallGemini(): único punto de entrada del chat del participante.",
            "BuildSystemPreamble(): reglas del “profesor” (no revelar respuesta correcta, pistas breves).",
            "apiTurns: memoria del chat solo dentro de la pregunta actual (máx. 30 turnos).",
            "ProcessResponse(): muestra respuesta y, si el avatar está activo, llama Speak().",
            "Manejo de errores de red (429, 503 con reintentos) y mensajes al participante.",
            "ClearChatHistory(): al pasar de pregunta, reinicia chat visible y contexto API.",
            "Exportación secundaria a Logs/ExperimentData_{UserID}.csv.",
        ],
        extra="Claves API: campo apiKey (placeholder YOUR_GEMINI_API_KEY_HERE en el repositorio).",
    )

    add_file_section(
        doc,
        "Scripts/Core/DataLogger.cs",
        "Registro principal de respuestas para análisis.",
        [
            "Escribe en CSV data/ExperimentData_{UserID}.csv (junto al ejecutable).",
            "Columnas: UserID, escenario, pregunta, respuesta, correcta, confianza, tiempo, timestamp.",
            "Bloqueo de archivo y reintentos si Excel tiene el CSV abierto.",
            "DeleteIncompleteFile(): elimina archivo si el participante sale antes de completar el primer escenario.",
        ],
    )

    add_file_section(
        doc,
        "Scripts/Core/AvatarDisplayController.cs",
        "Muestra el personaje 3D en el panel derecho (solo condición C).",
        [
            "Cámara secundaria y RenderTexture 1024×1024.",
            "Renderiza el modelo en capa AvatarLayer y lo muestra en RawImage del panel UI.",
            "Encuadra cuerpo completo y oculta el modelo de la vista principal mientras se muestra en el panel.",
            "Se activa cuando QuestionManager enciende characterModel (escenario índice 2).",
        ],
    )

    add_file_section(
        doc,
        "Scripts/Audio/AgentSpeechController.cs",
        "Punto de entrada para que el agente hable.",
        [
            "Método Speak(string text).",
            "Delega en AzureLipSync para síntesis de voz.",
            "Lo invoca ExperimentLogic cuando hay respuesta de Gemini y avatar visible.",
        ],
    )

    add_file_section(
        doc,
        "Scripts/Audio/AzureLipSync.cs",
        "Texto a voz con Azure Cognitive Services y sincronización labial.",
        [
            "Voz es-MX-JorgeNeural vía Microsoft.CognitiveServices.Speech.",
            "Convierte PCM a AudioClip y reproduce por AudioSource.",
            "Visemas → blend shapes de boca (A, E, U, mandíbula).",
            "Animaciones TTB_idle1 / TTB_talk2 (parámetro isTalking).",
            "Pausa audio si la ventana pierde foco; avisa en chat si falla TTS.",
        ],
        extra="Claves API: subscriptionKey y region (placeholder en repositorio).",
    )

    add_file_section(
        doc,
        "Editor/StandaloneBuild.cs",
        "Compilación automática del .exe desde línea de comandos (solo Editor Unity).",
        [
            "Método PerformWindowsBuild().",
            "Genera Build/Windows/ExperimentPrototypeB03230.exe.",
            "Invocado por _tools/build_windows.bat con -batchmode -executeMethod.",
        ],
        extra="Nota: envuelto en #if UNITY_EDITOR; no va en el build final.",
    )

    # --- Terceros ---
    doc.add_heading("Script del asset de terceros (Assets/1toonteen/)", level=1)

    add_file_section(
        doc,
        "Scripts/blendshapesshow.cs",
        "Utilidad de demostración del paquete OneToonTeen (no forma parte del flujo del experimento).",
        [
            "Recorre blend shapes de la cabeza uno por uno con animación.",
            "Al terminar cierra la aplicación.",
            "El experimento usa AzureLipSync para la boca en condición C, no este script.",
        ],
    )

    # --- Tools ---
    doc.add_heading("Herramientas auxiliares (_tools/)", level=1)
    add_body(doc, "Scripts fuera de Unity para mantenimiento, documentación y build.")

    add_table(
        doc,
        ["Archivo", "Qué hace"],
        [
            ["build_windows.bat", "Ejecuta Unity en batch y compila el .exe. Requiere Unity 6000.3.11f1 y Editor cerrado."],
            ["clean_for_delivery.py", "Limpia CSVs de prueba, Temp/obj; reemplaza claves API por placeholders antes de GitHub."],
            ["generate_scenarios_yaml.py", "Lee preguntas desde docx del curso e inyecta escenarios en SampleScene.unity."],
            ["generate_entregable_figures.py", "Genera figuras del Entregable 2 con matplotlib → docs/figures/."],
            ["generate_entregable2_docx.py", "Convierte el Entregable 2 de Markdown a Word/PDF con formato UCR."],
            ["generate_mapa_codigo_docx.py", "Genera este documento Word (MAPA_CODIGO.docx)."],
        ],
    )

    # --- No experiment code ---
    doc.add_heading("Archivos que no son código del experimento", level=1)
    add_table(
        doc,
        ["Ubicación", "Contenido"],
        [
            ["Assets/TextMesh Pro/Shaders/*.shader", "Shaders de TextMesh Pro (texto UI)."],
            ["Assets/Packages/", "DLLs NuGet: Azure Speech SDK, Azure.Core, System.Text.Json, etc."],
            ["Assets/1toonteen/ (excepto blendshapesshow.cs)", "Modelo 3D, materiales, animaciones del avatar."],
            ["Assets/NuGet/", "Configuración del gestor NuGet para Unity."],
            ["ProjectSettings/", "Resolución, pantalla completa, input — configuración del player."],
        ],
    )

    # --- Data paths ---
    doc.add_heading("Dónde se guardan los datos al ejecutar", level=1)
    add_table(
        doc,
        ["Ruta", "Origen", "Contenido"],
        [
            ["CSV data/ExperimentData_{ID}.csv", "DataLogger", "Principal — respuestas, confianza, tiempos"],
            ["Logs/ExperimentData_{ID}.csv", "ExperimentLogic", "Secundario — acciones y chat por pregunta"],
            ["Logs/standalone-build.log", "Build batch", "Log de compilación Unity"],
        ],
    )
    add_body(doc, "En Git, CSV data/ y Logs/ con datos reales están en .gitignore.")

    # --- Dependencies ---
    doc.add_heading("Dependencias entre scripts", level=1)
    add_code_block(
        doc,
        [
            "QuestionManager",
            "  ├── ExperimentLogic   (chat, consentimiento, Gemini)",
            "  ├── DataLogger        (CSV principal)",
            "  └── characterModel    (GameObject del avatar)",
            "",
            "ExperimentLogic",
            "  ├── QuestionManager   (input chat, escenario activo)",
            "  └── AgentSpeechController → AzureLipSync",
            "",
            "AvatarDisplayController",
            "  └── QuestionManager.characterModel + panel UI derecho",
        ],
    )

    # --- How to edit ---
    doc.add_heading("Cómo abrir y editar", level=1)
    add_numbered = [
        "Abrir el proyecto en Unity Hub → versión 6000.3.11f1.",
        "Escena: Assets/Scenes/SampleScene.unity.",
        "En la jerarquía: QuestionManager, ExperimentLogic, DataLogger, AvatarDisplayController, AzureLipSync.",
        "Preguntas: lista scenarios en QuestionManager (Inspector) o regenerar con generate_scenarios_yaml.py.",
        "Build: _tools/build_windows.bat o menú Build en Unity.",
    ]
    for i, step in enumerate(add_numbered, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Última actualización: junio 2026 — incluye guía detallada de chat, voz Azure y avatar (condiciones A/B/C)."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_MUTED

    return doc


def main() -> int:
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(DOCX_PATH)
    print(f"Generado: {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
